import zipfile
from datetime import datetime, timezone
from io import BytesIO

import pytest
from docx import Document

from app.schemas.interview_report import StudentFeedbackExportRequest
from app.services import admin_dashboard_service
from app.services.admin_dashboard_service import (
    _add_student_feedback_section,
    _export_filename_for_scope,
    _new_feedback_document,
    _sanitize_filename,
    group_reports_by_student_id,
    resolve_student_export_scope,
)


class _FakeAggregateCursor:
    def __init__(self, rows):
        self.rows = rows

    async def to_list(self, length=None):
        return list(self.rows)


class _FakeCollection:
    def __init__(self, rows):
        self.rows = rows

    def aggregate(self, pipeline):
        return _FakeAggregateCursor(self.rows)


def _mock_db_for_reports(monkeypatch, rows):
    fake_db = {"interview_reports": _FakeCollection(rows)}
    monkeypatch.setattr(admin_dashboard_service, "get_database", lambda: fake_db)
    return fake_db


def test_student_export_docx_structure_omits_marker_fields():
    report = {
        "student": {"name": "Suhas Pulapa"},
        "company": {"name": "Drip Off AI"},
        "opportunity": {"role": "Software Engineer Intern"},
        "session": {"scheduled_at": datetime(2026, 8, 6, tzinfo=timezone.utc)},
        "visible_to_student": True,
        "overall": {"score": 6, "summary": "Strong overall performance."},
        "interviewer_satisfaction": "Met the bar with depth.",
        "answers": [{
            "question_text": "Tell me about yourself.",
            "student_answer": "I built a product team.",
            "accuracy": 16,
            "ideal_answer": "Explain background and impact.",
            "feedback": "Clear and direct.",
        }],
        "strengths": ["Clear communication"],
        "improvements": ["Need more depth in system design"],
        "communication": {"notes": "Good communication and confidence."},
        "skill_ratings": {"python": 4},
        "coaching_note": "Practice mock interviews.",
    }

    document = _new_feedback_document("Student Interview Feedback")
    _add_student_feedback_section(document, report, 1)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Student 1" not in text
    assert "Status" not in text
    assert "Pending" not in text and "Published" not in text
    assert "COACHING FOR NEXT TIME" not in text
    assert "End of Student" not in text
    assert "student interview feedback" in text.lower()
    assert "overall feedback" in text.lower()
    assert "how they met the bar" in text.lower()
    assert "questions & answers" in text.lower()


def test_student_scope_uses_canonical_student_id_and_filters_history():
    reports = [
        {"_id": "r1", "student_id": "student-123", "company": {"name": "Drip Off AI"}},
        {"_id": "r2", "student_id": "student-123", "company": {"name": "Totem Interactive"}},
        {"_id": "r3", "student_id": "student-456", "company": {"name": "Other Company"}},
    ]

    resolved = resolve_student_export_scope(reports, scope="student", student_id="student-123")

    assert resolved["student_id"] == "student-123"
    assert [report["_id"] for report in resolved["reports"]] == ["r1", "r2"]


def test_student_scope_rejects_missing_student_id():
    reports = [{"_id": "r1", "student_id": None, "student": {"name": "Suhas Pulapa"}}]

    with pytest.raises(ValueError, match="student_id"):
        resolve_student_export_scope(reports, scope="student", student_id=None)


def test_student_export_request_accepts_student_ids():
    payload = StudentFeedbackExportRequest(
        report_ids=[],
        student_ids=["student-1", "student-2"],
        mode="combined",
        scope="selected",
    )

    assert payload.student_ids == ["student-1", "student-2"]


def test_single_scope_keeps_only_selected_report():
    reports = [
        {"_id": "r1", "student_id": "student-123", "company": {"name": "Drip Off AI"}},
        {"_id": "r2", "student_id": "student-123", "company": {"name": "Totem Interactive"}},
    ]

    resolved = resolve_student_export_scope(reports, scope="single", student_id="student-123")

    assert resolved["student_id"] == "student-123"
    assert [report["_id"] for report in resolved["reports"]] == ["r1"]


def test_group_reports_by_student_id_counts_unique_companies_and_reports():
    reports = [
        {"_id": "r1", "student_id": "student-123", "student": {"name": "Suhas Pulapa"}, "company_id": "company-1", "company": {"name": "Drip Off AI"}},
        {"_id": "r2", "student_id": "student-123", "student": {"name": "Suhas Pulapa"}, "company_id": "company-1", "company": {"name": "Drip Off AI"}},
        {"_id": "r3", "student_id": "student-123", "student": {"name": "Suhas Pulapa"}, "company_id": "company-2", "company": {"name": "Totem Interactive"}},
        {"_id": "r4", "student_id": "student-456", "student": {"name": "Mahesh Kumar"}, "company_id": "company-3", "company": {"name": "LocalButcher"}},
    ]

    grouped = group_reports_by_student_id(reports)

    assert len(grouped) == 2
    suhas = next(group for group in grouped if group["student_id"] == "student-123")
    assert suhas["student_name"] == "Suhas Pulapa"
    assert suhas["company_count"] == 2
    assert suhas["interview_count"] == 3
    assert [report["_id"] for report in suhas["reports"]] == ["r1", "r2", "r3"]


def test_sanitize_filename_removes_unsafe_chars():
    assert _sanitize_filename("Suhas Pulapa") == "Suhas_Pulapa"
    assert _sanitize_filename("Drip Off AI") == "Drip_Off_AI"
    assert _sanitize_filename("Suhas-Pulapa") == "Suhas_Pulapa"
    assert _sanitize_filename("Thakkuri Shiva Kumar") == "Thakkuri_Shiva_Kumar"


def test_export_filename_for_scope_single_interview():
    """Single company/interview should include the company name."""
    filename = _export_filename_for_scope(
        "Suhas Pulapa",
        ["Drip Off AI"],
        scope="single",
        total_student_companies=4,
    )
    assert filename == "Suhas_Pulapa_Drip_Off_AI_Interview_Feedback.docx"


def test_export_filename_for_scope_all_companies():
    """All companies for student should NOT include company names."""
    filename = _export_filename_for_scope(
        "Suhas Pulapa",
        ["Drip Off AI", "Totem Interactive", "LocalButcher", "GTM / Gen AI"],
        scope="student",
        total_student_companies=4,
    )
    assert filename == "Suhas_Pulapa_Interview_Feedback.docx"


def test_export_filename_for_scope_selected_companies():
    """Selected subset of companies (not all) should include 'Selected_Companies'."""
    filename = _export_filename_for_scope(
        "Suhas Pulapa",
        ["Drip Off AI", "LocalButcher"],
        scope="selected",
        total_student_companies=4,
    )
    assert filename == "Suhas_Pulapa_Selected_Companies_Interview_Feedback.docx"


def test_export_filename_for_scope_selected_that_equals_all():
    """Selected scope that happens to equal all companies (edge case)."""
    filename = _export_filename_for_scope(
        "Suhas Pulapa",
        ["Drip Off AI"],
        scope="selected",
        total_student_companies=1,
    )
    assert filename == "Suhas_Pulapa_Interview_Feedback.docx"


@pytest.mark.parametrize(
    ("selected_count", "mode", "expected_names"),
    [
        (1, "combined", {"Student_Interview_Feedback.docx"}),
        (2, "combined", {"Student_Interview_Feedback_Combined.docx"}),
        (3, "combined", {"Student_Interview_Feedback_Combined.docx"}),
        (4, "combined", {"Student_Interview_Feedback_Combined.docx"}),
        (5, "combined", {"Student_Interview_Feedback_Combined.docx"}),
        (1, "separate", {"Alpha_Company_1_Interview_Feedback.docx"}),
        (2, "separate", {"Alpha_Interview_Feedback.docx", "Bravo_Interview_Feedback.docx"}),
        (4, "separate", {"Alpha_Interview_Feedback.docx", "Bravo_Interview_Feedback.docx", "Charlie_Interview_Feedback.docx", "Delta_Interview_Feedback.docx"}),
        (5, "separate", {"Alpha_Interview_Feedback.docx", "Bravo_Interview_Feedback.docx", "Charlie_Interview_Feedback.docx", "Delta_Interview_Feedback.docx", "Echo_Interview_Feedback.docx"}),
        (2, "both", {"Student_Interview_Feedback_Combined.docx", "Alpha_Interview_Feedback.docx", "Bravo_Interview_Feedback.docx"}),
        (5, "both", {"Student_Interview_Feedback_Combined.docx", "Alpha_Interview_Feedback.docx", "Bravo_Interview_Feedback.docx", "Charlie_Interview_Feedback.docx", "Delta_Interview_Feedback.docx", "Echo_Interview_Feedback.docx"}),
    ],
)
@pytest.mark.asyncio
async def test_bulk_student_exports_by_selected_count(monkeypatch, selected_count, mode, expected_names):
    student_names = ["Alpha", "Bravo", "Charlie", "Delta", "Echo"]
    reports = []
    for index, name in enumerate(student_names[:selected_count], start=1):
        reports.append(
            {
                "_id": f"64d0f71e1a2b3c4d5e6f78{index:02d}",
                "student_id": f"student-{index}",
                "student": {"name": name},
                "company_id": f"company-{index}",
                "company": {"name": f"Company {index}"},
                "opportunity": {"role": "Frontend Engineer"},
                "session": {"scheduled_at": datetime(2026, 8, index, tzinfo=timezone.utc)},
                "overall": {"score": 7 + index, "summary": "Strong"},
                "answers": [],
            }
        )

    _mock_db_for_reports(monkeypatch, reports)
    from app.services.admin_dashboard_service import build_student_feedback_export

    content, filename, media_type = await build_student_feedback_export(
        [report["_id"] for report in reports],
        mode,
        scope="selected",
    )

    if mode == "combined":
        assert filename == "Student_Interview_Feedback_Combined.docx" if selected_count > 1 else "Student_Interview_Feedback.docx"
        assert media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        expected_zip_name = "Student_Interview_Feedback.zip" if selected_count > 1 else f"{student_names[0]}_Interview_Feedback.zip"
        assert filename == expected_zip_name
        assert media_type == "application/zip"
        with zipfile.ZipFile(BytesIO(content)) as archive:
            assert set(archive.namelist()) == expected_names


@pytest.mark.asyncio
async def test_multi_student_separate_export_zips_student_files(monkeypatch):
    reports = [
        {
            "_id": "64d0f71e1a2b3c4d5e6f7890",
            "student_id": "student-1",
            "student": {"name": "Bharath"},
            "company_id": "company-1",
            "company": {"name": "Drip Off AI"},
            "opportunity": {"role": "Frontend Engineer"},
            "session": {"scheduled_at": datetime(2026, 8, 6, tzinfo=timezone.utc)},
            "overall": {"score": 7, "summary": "Strong"},
            "answers": [],
        },
        {
            "_id": "64d0f71e1a2b3c4d5e6f7891",
            "student_id": "student-2",
            "student": {"name": "Bhargav Martha"},
            "company_id": "company-2",
            "company": {"name": "Totem Interactive"},
            "opportunity": {"role": "Data Analyst"},
            "session": {"scheduled_at": datetime(2026, 8, 7, tzinfo=timezone.utc)},
            "overall": {"score": 8, "summary": "Strong"},
            "answers": [],
        },
    ]

    _mock_db_for_reports(monkeypatch, reports)
    from app.services.admin_dashboard_service import build_student_feedback_export

    content, filename, media_type = await build_student_feedback_export(
        ["64d0f71e1a2b3c4d5e6f7890", "64d0f71e1a2b3c4d5e6f7891"],
        "separate",
        scope="selected",
    )

    assert filename == "Student_Interview_Feedback.zip"
    assert media_type == "application/zip"
    with zipfile.ZipFile(BytesIO(content)) as archive:
        assert set(archive.namelist()) == {"Bharath_Interview_Feedback.docx", "Bhargav_Martha_Interview_Feedback.docx"}


@pytest.mark.asyncio
async def test_multi_student_both_export_includes_combined_doc_and_student_files(monkeypatch):
    reports = [
        {
            "_id": "64d0f71e1a2b3c4d5e6f7892",
            "student_id": "student-1",
            "student": {"name": "Bharath"},
            "company_id": "company-1",
            "company": {"name": "Drip Off AI"},
            "opportunity": {"role": "Frontend Engineer"},
            "session": {"scheduled_at": datetime(2026, 8, 6, tzinfo=timezone.utc)},
            "overall": {"score": 7, "summary": "Strong"},
            "answers": [],
        },
        {
            "_id": "64d0f71e1a2b3c4d5e6f7893",
            "student_id": "student-2",
            "student": {"name": "Bhargav Martha"},
            "company_id": "company-2",
            "company": {"name": "Totem Interactive"},
            "opportunity": {"role": "Data Analyst"},
            "session": {"scheduled_at": datetime(2026, 8, 7, tzinfo=timezone.utc)},
            "overall": {"score": 8, "summary": "Strong"},
            "answers": [],
        },
    ]

    _mock_db_for_reports(monkeypatch, reports)
    from app.services.admin_dashboard_service import build_student_feedback_export

    content, filename, media_type = await build_student_feedback_export(
        ["64d0f71e1a2b3c4d5e6f7892", "64d0f71e1a2b3c4d5e6f7893"],
        "both",
        scope="selected",
    )

    assert filename == "Student_Interview_Feedback.zip"
    assert media_type == "application/zip"
    with zipfile.ZipFile(BytesIO(content)) as archive:
        assert set(archive.namelist()) == {
            "Student_Interview_Feedback_Combined.docx",
            "Bharath_Interview_Feedback.docx",
            "Bhargav_Martha_Interview_Feedback.docx",
        }


@pytest.mark.asyncio
async def test_combined_student_export_has_single_title_and_company_heading(monkeypatch):
    report = {
        "_id": "64d0f71e1a2b3c4d5e6f7894",
        "student_id": "student-10",
        "student": {"name": "Chaitanya Jyothi"},
        "company_id": "company-10",
        "company": {"name": "Nexuses"},
        "opportunity": {"role": "Software Engineer"},
        "session": {"scheduled_at": datetime(2026, 8, 6, tzinfo=timezone.utc)},
        "visible_to_student": True,
        "overall": {"score": 8, "summary": "Strong overall performance."},
        "interviewer_satisfaction": "Good depth and clarity.",
        "answers": [{
            "question_text": "Tell me about yourself.",
            "student_answer": "I built a product team.",
            "accuracy": 16,
            "ideal_answer": "Explain background and impact.",
            "feedback": "Clear and direct.",
        }],
        "strengths": ["Clear communication"],
        "improvements": ["Need more depth in system design"],
        "communication": {"notes": "Good communication and confidence."},
        "skill_ratings": {"python": 4},
    }

    _mock_db_for_reports(monkeypatch, [report])
    from app.services.admin_dashboard_service import build_student_feedback_export

    content, filename, media_type = await build_student_feedback_export([report["_id"]], "combined", scope="selected")

    assert filename == "Chaitanya_Jyothi_Interview_Feedback.docx"
    assert media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    title_count = sum(1 for text in paragraphs if text == "STUDENT INTERVIEW FEEDBACK")
    assert title_count == 1
    assert "Student Interview Feedback" not in paragraphs
    assert "Student: Chaitanya Jyothi" in paragraphs
    assert "COMPANIES INTERVIEWED" in paragraphs
    assert "1. Nexuses" in paragraphs

