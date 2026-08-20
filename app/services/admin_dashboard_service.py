import asyncio
import re
import zipfile
from io import BytesIO
from datetime import datetime, timedelta, timezone

from fastapi import HTTPException, status
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from app.db.collections import (
    APPLICATIONS,
    COMPANIES,
    HIRING_OPPORTUNITIES,
    INTERVIEW_REPORTS,
    INTERVIEW_SESSIONS,
    STUDENTS,
    TRANSCRIPTS,
)
from app.db.mongodb import get_database
from app.models.application import normalize_application_status
from app.services.transcript_service import resolve_interview_date
from app.utils.mongo import serialize_mongo
from app.utils.object_id import to_object_id


def _sort_key_for_date(val: object) -> datetime:
    """Normalize various date representations into a UTC datetime for sorting and comparisons."""
    if isinstance(val, datetime):
        if val.tzinfo is None:
            return val.replace(tzinfo=timezone.utc)
        return val.astimezone(timezone.utc)
    if isinstance(val, str) and val.strip():
        try:
            dt = datetime.fromisoformat(val.replace("Z", "+00:00"))
            if dt.tzinfo is None:
                return dt.replace(tzinfo=timezone.utc)
            return dt.astimezone(timezone.utc)
        except ValueError:
            pass
    return datetime.min.replace(tzinfo=timezone.utc)


REAL_APPLICATION_FILTER = {
    "$or": [
        {"application_details.interested": {"$exists": True, "$ne": False}},
        {
            "application_details": {"$exists": False},
            "is_interested": {"$ne": False},
            "status": {"$ne": "not_interested"},
        },
    ]
}
NOT_INTERESTED_FILTER = {
    "$or": [
        {"application_details.interested": False},
        {"application_details": {"$exists": False}, "is_interested": False},
        {"application_details": {"$exists": False}, "status": "not_interested"},
    ]
}


async def get_admin_dashboard() -> dict:
    db = get_database()

    # These are all independent, so fire them concurrently (one round-trip wave
    # instead of nine sequential ones — Atlas latency dominated the old path).
    (
        total_students, total_companies, total_opportunities, response_count,
        total_applications, not_interested_count, shortlisted_count, rejected_count, hired_count,
    ) = await asyncio.gather(
        db[STUDENTS].count_documents({}),
        db[COMPANIES].count_documents({}),
        db[HIRING_OPPORTUNITIES].count_documents({}),
        db[APPLICATIONS].count_documents({}),
        db[APPLICATIONS].count_documents(REAL_APPLICATION_FILTER),
        db[APPLICATIONS].count_documents(NOT_INTERESTED_FILTER),
        db[APPLICATIONS].count_documents(
            {"$or": [{"current_status": "SHORTLISTED"}, {"current_status": {"$exists": False}, "status": "shortlisted"}]}
        ),
        db[APPLICATIONS].count_documents(
            {"$or": [{"current_status": "REJECTED"}, {"current_status": {"$exists": False}, "status": "rejected"}]}
        ),
        db[APPLICATIONS].count_documents(
            {"$or": [{"current_status": {"$in": ["SELECTED", "JOINED"]}}, {"current_status": {"$exists": False}, "status": "hired"}]}
        ),
    )

    status_task = db[APPLICATIONS].aggregate(
        [
            {"$group": {"_id": {"$ifNull": ["$current_status", "$status"]}, "count": {"$sum": 1}}},
            {"$sort": {"count": -1, "_id": 1}},
        ]
    ).to_list(length=None)

    recent_apps_task = list_recent_applications(limit=8)
    opportunity_pipeline = [
        {"$match": {"deleted_at": {"$exists": False}}},
        {
            "$lookup": {
                "from": APPLICATIONS,
                "localField": "_id",
                "foreignField": "opportunity_id",
                "as": "applications",
            }
        },
        {
            "$lookup": {
                "from": COMPANIES,
                "localField": "company_id",
                "foreignField": "_id",
                "as": "company",
            }
        },
        {"$unwind": {"path": "$company", "preserveNullAndEmptyArrays": True}},
        {
            "$project": {
                "role": 1,
                "tech_stack": 1,
                "must_have_skills": 1,
                "location": 1,
                "stipend": 1,
                "duration": 1,
                "company_status": 1,
                "opportunity_received_at": 1,
                "company": {"_id": "$company._id", "name": "$company.name"},
                "response_count": {"$size": "$applications"},
                "application_count": {
                    "$size": {
                        "$filter": {
                            "input": "$applications",
                            "as": "application",
                            "cond": {
                                "$and": [
                                    {
                                        "$ne": [
                                            {
                                                "$ifNull": [
                                                    "$$application.application_details.interested",
                                                    "$$application.is_interested",
                                                ]
                                            },
                                            False,
                                        ]
                                    },
                                    {"$ne": [{"$ifNull": ["$$application.current_status", "$$application.status"]}, "not_interested"]},
                                ]
                            },
                        }
                    }
                },
                "shortlists_count": {"$ifNull": ["$shortlists_count", 0]},
            }
        },
        {"$sort": {"opportunity_received_at": -1, "updated_at": -1}},
        {"$limit": 500},
    ]
    recent_opps_task = db[HIRING_OPPORTUNITIES].aggregate(opportunity_pipeline).to_list(length=None)

    repeated_task = db[HIRING_OPPORTUNITIES].aggregate(
        [
            {
                "$group": {
                    "_id": "$company_id",
                    "opportunity_count": {"$sum": 1},
                    "roles": {"$addToSet": "$role"},
                    "last_received_at": {"$max": "$opportunity_received_at"},
                }
            },
            {"$match": {"opportunity_count": {"$gt": 1}}},
            {"$lookup": {"from": COMPANIES, "localField": "_id", "foreignField": "_id", "as": "company"}},
            {"$unwind": "$company"},
            {
                "$project": {
                    "company": {"_id": "$company._id", "name": "$company.name"},
                    "opportunity_count": 1,
                    "roles": 1,
                    "last_received_at": 1,
                }
            },
            {"$sort": {"opportunity_count": -1, "last_received_at": -1}},
            {"$limit": 8},
        ]
    ).to_list(length=None)

    # Run the four aggregations concurrently rather than one after another.
    status_breakdown, recent_applications, recent_opportunities, repeated_companies = await asyncio.gather(
        status_task, recent_apps_task, recent_opps_task, repeated_task
    )

    # ---- Overview funnel + action center: fire every read concurrently ----
    cutoff = datetime.now(timezone.utc) - timedelta(days=30)
    (
        interviewing_count, dropped_count, awaiting_count, not_shortlisted_count,
        missing_shortlist, profiles_pending, reports_total, reports_published,
        report_app_ids_raw, active_recent_list, ever_applied_list, questions_banked,
        placed_students,
    ) = await asyncio.gather(
        db[APPLICATIONS].count_documents(
            {"current_status": {"$in": ["INTERVIEW_IN_PROGRESS", "INTERVIEW_SCHEDULED", "INTERVIEW_COMPLETED"]}}
        ),
        db[APPLICATIONS].count_documents({"current_status": "DROPPED"}),
        db[APPLICATIONS].count_documents({"current_status": "APPLIED", "screening.decision": {"$in": [None, "none"]}}),
        db[APPLICATIONS].count_documents(
            {"$or": [
                {"current_status": "NOT_SHORTLISTED"},
                {"screening.decision": {"$in": ["not_shortlisted", "selected_elsewhere", "resume_not_found"]}},
            ]}
        ),
        db[HIRING_OPPORTUNITIES].count_documents(
            {"shortlist_imported_at": {"$in": [None]}, "responses_imported_at": {"$ne": None}}
        ),
        db[HIRING_OPPORTUNITIES].count_documents(
            {
                "profiles_requested": {"$nin": [None, "", 0, "0"]},
                "$or": [{"profiles_shared": {"$in": [None, "", 0, "0"]}}, {"profiles_shared": {"$exists": False}}],
            }
        ),
        db[INTERVIEW_REPORTS].count_documents({}),
        db[INTERVIEW_REPORTS].count_documents({"visible_to_student": True}),
        db[INTERVIEW_REPORTS].distinct("application_id"),
        db[APPLICATIONS].distinct("student_id", {"applied_at": {"$gte": cutoff}}),
        db[APPLICATIONS].distinct("student_id"),
        db["questions"].count_documents({}),
        db[APPLICATIONS].distinct("student_id", {"current_status": {"$in": ["SELECTED", "JOINED", "OFFER_ACCEPTED"]}}),
    )
    report_app_ids = [i for i in report_app_ids_raw if i]
    interviewed_no_report = await db[APPLICATIONS].count_documents(
        {"current_status": "INTERVIEW_IN_PROGRESS", "_id": {"$nin": report_app_ids}}
    )
    inactive_30d = len(set(ever_applied_list) - set(active_recent_list))
    placed = len(placed_students)

    funnel = [
        {"key": "applied", "label": "Applied", "sub": "all applications", "n": response_count},
        {"key": "interested", "label": "Interested", "sub": "student opted in", "n": total_applications},
        {"key": "shortlisted", "label": "Shortlisted", "sub": "company picked them", "n": shortlisted_count},
        {"key": "interviewing", "label": "Interviewing", "sub": "in process now", "n": interviewing_count},
        {"key": "placed", "label": "Selected / joined", "sub": "placed", "n": hired_count},
    ]

    action_center = {
        "missing_shortlist_data": missing_shortlist,
        "profiles_requested_not_shared": profiles_pending,
        "reports_unpublished": reports_total - reports_published,
        "interviewed_no_report": interviewed_no_report,
        "inactive_30d": inactive_30d,
    }
    action_total = sum(action_center.values())

    return serialize_mongo(
        {
            "summary": {
                "total_students": total_students,
                "total_companies": total_companies,
                "total_opportunities": total_opportunities,
                "response_count": response_count,
                "total_applications": total_applications,
                "not_interested_count": not_interested_count,
                "shortlisted_count": shortlisted_count,
                "interview_ready_count": shortlisted_count,
                "rejected_count": rejected_count,
                "hired_count": hired_count,
            },
            "funnel": funnel,
            "loss": {"dropped": dropped_count, "awaiting": awaiting_count, "not_shortlisted": not_shortlisted_count},
            "action_center": action_center,
            "action_total": action_total,
            "placement": {"placed": placed, "total_students": total_students,
                          "rate": round(placed / total_students * 100, 1) if total_students else 0},
            "reports_summary": {"reports": reports_total, "published": reports_published,
                                "pending": reports_total - reports_published, "questions": questions_banked},
            "status_breakdown": [{"status": item["_id"] or "unknown", "count": item["count"]} for item in status_breakdown],
            "recent_applications": recent_applications,
            "recent_opportunities": recent_opportunities,
            "repeated_companies": repeated_companies,
        }
    )


async def list_recent_applications(limit: int = 50, status_value: str | None = None) -> list[dict]:
    db = get_database()
    match_stage = dict(REAL_APPLICATION_FILTER)
    if status_value:
        normalized_status = normalize_application_status(status_value)
        match_stage["$and"] = [
            {
                "$or": [
                    {"current_status": normalized_status},
                    {"current_status": {"$exists": False}, "status": status_value},
                ]
            }
        ]

    pipeline = [
        {"$match": match_stage},
        {"$sort": {"applied_at": -1, "created_at": -1}},
        {"$limit": limit},
        {"$lookup": {"from": STUDENTS, "localField": "student_id", "foreignField": "_id", "as": "student"}},
        {"$unwind": {"path": "$student", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "company"}},
        {"$unwind": {"path": "$company", "preserveNullAndEmptyArrays": True}},
        {
            "$lookup": {
                "from": HIRING_OPPORTUNITIES,
                "localField": "opportunity_id",
                "foreignField": "_id",
                "as": "opportunity",
            }
        },
        {"$unwind": {"path": "$opportunity", "preserveNullAndEmptyArrays": True}},
        {
            "$project": {
                "current_status": 1,
                "final_status": 1,
                "status": {"$ifNull": ["$current_status", "$status"]},
                "is_interested": {"$ifNull": ["$application_details.interested", "$is_interested"]},
                "applied_at": 1,
                "application_details": 1,
                "github_link": {"$ifNull": ["$application_details.github_link", "$github_link"]},
                "project_link": {"$ifNull": ["$application_details.project_link", "$project_link"]},
                "resume_link": {"$ifNull": ["$application_details.submitted_resume_url", "$resume_link"]},
                "student": {
                    "_id": "$student._id",
                    "name": "$student.name",
                    "email": "$student.email",
                    "phone": "$student.phone",
                    "college": "$student.college",
                    "degree": "$student.degree",
                    "department": "$student.department",
                    "year_of_passing": "$student.year_of_passing",
                },
                "company": {"_id": "$company._id", "name": "$company.name"},
                "opportunity": {
                    "_id": "$opportunity._id",
                    "role": "$opportunity.role",
                    "tech_stack": "$opportunity.tech_stack",
                    "must_have_skills": "$opportunity.must_have_skills",
                    "location": "$opportunity.location",
                    "stipend": "$opportunity.stipend",
                    "duration": "$opportunity.duration",
                    "opportunity_received_at": "$opportunity.opportunity_received_at",
                },
            }
        },
    ]
    applications = await db[APPLICATIONS].aggregate(pipeline).to_list(length=limit)
    return serialize_mongo(applications)


async def list_admin_students(limit: int = 500) -> list[dict]:
    db = get_database()
    students = await db[STUDENTS].find({}).sort("name", 1).limit(limit).to_list(length=limit)
    student_ids = [student["_id"] for student in students]
    if not student_ids:
        return []

    application_pipeline = [
        {"$match": {"student_id": {"$in": student_ids}, **REAL_APPLICATION_FILTER}},
        {"$sort": {"applied_at": -1, "created_at": -1}},
        {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "company"}},
        {"$unwind": {"path": "$company", "preserveNullAndEmptyArrays": True}},
        {
            "$lookup": {
                "from": HIRING_OPPORTUNITIES,
                "localField": "opportunity_id",
                "foreignField": "_id",
                "as": "opportunity",
            }
        },
        {"$unwind": {"path": "$opportunity", "preserveNullAndEmptyArrays": True}},
        {
            "$project": {
                "student_id": 1,
                "current_status": 1,
                "final_status": 1,
                "status": {"$ifNull": ["$current_status", "$status"]},
                "applied_at": 1,
                "resume_link": {"$ifNull": ["$application_details.submitted_resume_url", "$resume_link"]},
                "github_link": {"$ifNull": ["$application_details.github_link", "$github_link"]},
                "project_link": {"$ifNull": ["$application_details.project_link", "$project_link"]},
                "company": {"_id": "$company._id", "name": "$company.name"},
                "opportunity": {
                    "_id": "$opportunity._id",
                    "role": "$opportunity.role",
                    "tech_stack": "$opportunity.tech_stack",
                    "must_have_skills": "$opportunity.must_have_skills",
                    "location": "$opportunity.location",
                    "opportunity_received_at": "$opportunity.opportunity_received_at",
                },
            }
        },
    ]
    applications = await db[APPLICATIONS].aggregate(application_pipeline).to_list(length=None)

    grouped: dict[str, list[dict]] = {}
    for application in applications:
        grouped.setdefault(str(application["student_id"]), []).append(application)

    student_rows = []
    for student in students:
        student_applications = grouped.get(str(student["_id"]), [])
        shortlisted = [item for item in student_applications if item.get("status") in {"SHORTLISTED", "shortlisted"}]
        not_shortlisted = [item for item in student_applications if item.get("status") not in {"SHORTLISTED", "shortlisted"}]
        student_rows.append(
            {
                "_id": student["_id"],
                "name": student.get("name"),
                "email": student.get("email"),
                "phone": student.get("phone"),
                "college": student.get("college"),
                "degree": student.get("degree"),
                "department": student.get("department"),
                "year_of_passing": student.get("year_of_passing"),
                "placed_status": bool(student.get("placed_status")),
                "application_count": len(student_applications),
                "shortlisted_count": len(shortlisted),
                "not_shortlisted_count": len(not_shortlisted),
                "applications": student_applications,
                "shortlisted_applications": shortlisted,
                "not_shortlisted_applications": not_shortlisted,
            }
        )

    return serialize_mongo(student_rows)


async def update_student_placement(student_id: str, placed_status: bool) -> dict:
    """Record the student's overall placement outcome from the admin dashboard."""
    db = get_database()
    try:
        object_id = to_object_id(student_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid student id")

    now = datetime.now(timezone.utc)
    result = await db[STUDENTS].update_one(
        {"_id": object_id},
        {"$set": {"placed_status": placed_status, "updated_at": now}},
    )
    if not result.matched_count:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Student not found")

    return serialize_mongo(await db[STUDENTS].find_one({"_id": object_id}))


SKILL_KEYS = ("python", "nodejs", "react", "mongodb", "sql", "dsa", "javascript")


# status expressions reused across aggregations
_STATUS = {"$ifNull": ["$current_status", "$status"]}
_INTERESTED = {"$ifNull": ["$application_details.interested", "$is_interested"]}
_SHORTLISTED_SET = ["SHORTLISTED", "shortlisted"]
_SELECTED_SET = ["SELECTED", "JOINED", "hired"]


def categorize_role(role: str | None, skills: str | None = None) -> str:
    """Bucket an opening by role/skills text. AI is checked first so an AI-flavoured
    full-stack role lands in AI, which is the split the dashboard cares about."""
    text = f"{role or ''} | {skills or ''}".lower()

    def has(*words: str) -> bool:
        return any(re.search(r"(?<![a-z])" + re.escape(w) + r"(?![a-z])", text) for w in words)

    if has("ai", "ml", "genai", "gen ai", "ai/ml", "llm", "rag", "agentic", "nlp") or (
        "machine learning" in text or "artificial intelligence" in text or "data scien" in text or "computer vision" in text
    ):
        return "AI / ML"
    if has("mern") or ("full stack" in text or "fullstack" in text or "full-stack" in text) or (has("react") and has("node")):
        return "MERN / Full Stack"
    if has("frontend", "react", "nextjs") or ("front end" in text or "front-end" in text or "next.js" in text):
        return "Frontend"
    if has("python", "django", "flask", "fastapi", "backend", "node") or ("back end" in text or "node.js" in text):
        return "Python / Backend"
    return "Other"


def _parse_date(value: str | None) -> datetime | None:
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M"):
        try:
            return datetime.strptime(value, fmt).replace(tzinfo=timezone.utc)
        except ValueError:
            continue
    return None


def _resolve_range(start: str | None, end: str | None) -> tuple[datetime, datetime]:
    now = datetime.now(timezone.utc)
    end_dt = _parse_date(end)
    end_dt = (end_dt + timedelta(days=1)) if end_dt else now  # end date is inclusive
    start_dt = _parse_date(start) or end_dt.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    return start_dt, end_dt


async def get_admin_analytics(start: str | None = None, end: str | None = None) -> dict:
    """Analytics for a date window (defaults to the current month), read off the
    normalized applications schema. Every metric below respects [start, end)."""
    db = get_database()
    start_dt, end_dt = _resolve_range(start, end)
    in_range = {"applied_at": {"$gte": start_dt, "$lt": end_dt}}

    totals = {
        "students": await db[STUDENTS].count_documents({}),
        "companies": await db[COMPANIES].count_documents({}),
        "opportunities": await db[HIRING_OPPORTUNITIES].count_documents({}),
        "applications": await db[APPLICATIONS].count_documents({}),
    }

    # ---- headline KPIs for the window (one pass) ----
    kpi_rows = await db[APPLICATIONS].aggregate(
        [
            {"$match": in_range},
            {
                "$group": {
                    "_id": None,
                    "applications": {"$sum": 1},
                    "students": {"$addToSet": "$student_id"},
                    "companies": {"$addToSet": "$company_id"},
                    "opportunities": {"$addToSet": "$opportunity_id"},
                    "interested": {"$sum": {"$cond": [{"$ne": [_INTERESTED, False]}, 1, 0]}},
                    "shortlisted": {"$sum": {"$cond": [{"$in": [_STATUS, _SHORTLISTED_SET]}, 1, 0]}},
                    "selected": {"$sum": {"$cond": [{"$in": [_STATUS, _SELECTED_SET]}, 1, 0]}},
                }
            },
        ]
    ).to_list(length=1)
    row = kpi_rows[0] if kpi_rows else {}
    applications = row.get("applications", 0)
    interested = row.get("interested", 0)
    shortlisted = row.get("shortlisted", 0)
    kpis = {
        "applications": applications,
        "students": len(row.get("students", [])),
        "companies": len(row.get("companies", [])),
        "opportunities": len(row.get("opportunities", [])),
        "interested": interested,
        "shortlisted": shortlisted,
        "selected": row.get("selected", 0),
        "interest_rate": round(interested / applications * 100, 1) if applications else 0,
        "shortlist_rate": round(shortlisted / interested * 100, 1) if interested else 0,
        "new_students": await db[STUDENTS].count_documents({"created_at": {"$gte": start_dt, "$lt": end_dt}}),
        "new_opportunities": await db[HIRING_OPPORTUNITIES].count_documents(
            {"opportunity_received_at": {"$gte": start_dt, "$lt": end_dt}}
        ),
    }

    # ---- daily trend (applications + distinct students per day) ----
    daily_raw = await db[APPLICATIONS].aggregate(
        [
            {"$match": in_range},
            {
                "$group": {
                    "_id": {"$dateToString": {"format": "%Y-%m-%d", "date": "$applied_at"}},
                    "apps": {"$sum": 1},
                    "students": {"$addToSet": "$student_id"},
                }
            },
            {"$project": {"_id": 0, "date": "$_id", "apps": 1, "students": {"$size": "$students"}}},
            {"$sort": {"date": 1}},
        ]
    ).to_list(length=None)

    # ---- status breakdown in window ----
    status_raw = await db[APPLICATIONS].aggregate(
        [
            {"$match": in_range},
            {"$group": {"_id": _STATUS, "n": {"$sum": 1}}},
            {"$sort": {"n": -1, "_id": 1}},
        ]
    ).to_list(length=None)
    status = [{"key": r["_id"] or "UNKNOWN", "n": r["n"]} for r in status_raw]

    # ---- per-student: how many opportunities each applied to ----
    per_student = await db[APPLICATIONS].aggregate(
        [
            {"$match": in_range},
            {
                "$group": {
                    "_id": "$student_id",
                    "apps": {"$sum": 1},
                    "shortlisted": {"$sum": {"$cond": [{"$in": [_STATUS, _SHORTLISTED_SET]}, 1, 0]}},
                }
            },
        ]
    ).to_list(length=None)

    buckets = {"1": 0, "2": 0, "3": 0, "4": 0, "5+": 0}
    for entry in per_student:
        n = entry["apps"]
        buckets["5+" if n >= 5 else str(n)] += 1
    apps_per_student = {
        "buckets": [{"label": label, "n": count} for label, count in buckets.items()],
        "avg": round(applications / len(per_student), 2) if per_student else 0,
        "max": max((e["apps"] for e in per_student), default=0),
    }

    # Every student who applied to more than one opening in the window (not a
    # top-N cap), most-active first. This is the admin's full activity roster.
    active = sorted([e for e in per_student if e["apps"] > 1], key=lambda e: (-e["apps"], -e["shortlisted"]))
    active_oids = [e["_id"] for e in active]
    id_to_name: dict = {}
    apps_by_student: dict = {}
    not_shortlisted_by_student: dict = {}
    if active_oids:
        students_docs = await db[STUDENTS].find(
            {"_id": {"$in": active_oids}},
            {"name": 1, "email": 1, "phone": 1, "external_user_id": 1},
        ).to_list(length=None)
        id_to_name = {s["_id"]: s for s in students_docs}
        app_rows = await db[APPLICATIONS].aggregate(
            [
                {"$match": {**in_range, "student_id": {"$in": active_oids}}},
                {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "opportunity_id", "foreignField": "_id", "as": "opp"}},
                {"$unwind": {"path": "$opp", "preserveNullAndEmptyArrays": True}},
                {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "co"}},
                {"$unwind": {"path": "$co", "preserveNullAndEmptyArrays": True}},
                {"$project": {
                    "student_id": 1, "applied_at": 1, "status": _STATUS,
                    "screening_remark": "$screening.remark",
                    "screening_decision": "$screening.decision",
                    "shortlist_done": {"$cond": [{"$ifNull": ["$opp.shortlist_imported_at", False]}, True, False]},
                    "role": "$opp.role", "skills": "$opp.must_have_skills", "company": "$co.name",
                }},
                {"$sort": {"applied_at": -1}},
            ]
        ).to_list(length=None)
        for a in app_rows:
            raw_status = a.get("status") or "APPLIED"
            decision = a.get("screening_decision") or ""
            # Admin sees the full picture: waitlisted (a backup, hidden from the
            # student), an explicit resume-stage reject, or simply not appearing
            # on an imported shortlist -> all surfaced instead of a bland APPLIED.
            display = raw_status
            if raw_status in ("APPLIED", "PROFILE_SHARED"):
                if decision == "waitlisted":
                    display = "WAITLISTED"
                elif decision in ("not_shortlisted", "selected_elsewhere", "resume_not_found") or a.get("shortlist_done"):
                    display = "NOT_SHORTLISTED"
                    not_shortlisted_by_student[a["student_id"]] = not_shortlisted_by_student.get(a["student_id"], 0) + 1
            apps_by_student.setdefault(a["student_id"], []).append(
                {
                    "company": a.get("company") or "Unknown",
                    "role": a.get("role") or "—",
                    "category": categorize_role(a.get("role"), a.get("skills")),
                    "status": display,
                    "remark": a.get("screening_remark"),
                    "applied_at": a.get("applied_at"),
                }
            )
    top_students = [
        {
            "id": e["_id"],
            "name": (id_to_name.get(e["_id"], {}) or {}).get("name") or "Unknown",
            "external_user_id": (id_to_name.get(e["_id"], {}) or {}).get("external_user_id"),
            "email": (id_to_name.get(e["_id"], {}) or {}).get("email"),
            "phone": (id_to_name.get(e["_id"], {}) or {}).get("phone"),
            "apps": e["apps"],
            "shortlisted": e["shortlisted"],
            "not_shortlisted": not_shortlisted_by_student.get(e["_id"], 0),
            "applications": apps_by_student.get(e["_id"], []),
        }
        for e in active
    ]

    # ---- top companies by applications in window ----
    top_companies_raw = await db[APPLICATIONS].aggregate(
        [
            {"$match": in_range},
            {"$group": {"_id": "$company_id", "n": {"$sum": 1}}},
            {"$sort": {"n": -1}},
            {"$limit": 8},
            {"$lookup": {"from": COMPANIES, "localField": "_id", "foreignField": "_id", "as": "c"}},
            {"$unwind": {"path": "$c", "preserveNullAndEmptyArrays": True}},
            {"$project": {"_id": 0, "name": "$c.name", "n": 1}},
        ]
    ).to_list(length=None)
    top_companies = [{"name": r.get("name") or "Unknown", "n": r["n"]} for r in top_companies_raw]

    funnel = [
        {"key": "Applied", "n": applications},
        {"key": "Interested", "n": interested},
        {"key": "Shortlisted", "n": shortlisted},
        {"key": "Selected", "n": kpis["selected"]},
    ]

    # ---- openings in window grouped by role category (AI / MERN / Python / …) ----
    opp_cat_rows = await db[APPLICATIONS].aggregate(
        [
            {"$match": in_range},
            {
                "$group": {
                    "_id": "$opportunity_id",
                    "apps": {"$sum": 1},
                    "shortlisted": {"$sum": {"$cond": [{"$in": [_STATUS, _SHORTLISTED_SET]}, 1, 0]}},
                    "company_id": {"$first": "$company_id"},
                }
            },
            {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "_id", "foreignField": "_id", "as": "opp"}},
            {"$unwind": {"path": "$opp", "preserveNullAndEmptyArrays": True}},
            {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "co"}},
            {"$unwind": {"path": "$co", "preserveNullAndEmptyArrays": True}},
            {"$project": {"apps": 1, "shortlisted": 1, "role": "$opp.role", "skills": "$opp.must_have_skills", "company": "$co.name"}},
        ]
    ).to_list(length=None)
    cats: dict = {}
    for r in opp_cat_rows:
        cat = categorize_role(r.get("role"), r.get("skills"))
        bucket = cats.setdefault(cat, {"category": cat, "opportunities": 0, "applications": 0, "shortlisted": 0, "companies": []})
        bucket["opportunities"] += 1
        bucket["applications"] += r.get("apps", 0)
        bucket["shortlisted"] += r.get("shortlisted", 0)
        bucket["companies"].append(
            {"company": r.get("company") or "Unknown", "role": r.get("role") or "—", "apps": r.get("apps", 0), "shortlisted": r.get("shortlisted", 0)}
        )
    for bucket in cats.values():
        bucket["companies"].sort(key=lambda x: -x["apps"])
    role_categories = sorted(cats.values(), key=lambda x: -x["opportunities"])

    # ---- all-time monthly context (unfiltered), for the trend backdrop ----
    by_month_raw = await db[APPLICATIONS].aggregate(
        [
            {"$match": {"applied_at": {"$ne": None}}},
            {"$group": {"_id": {"$dateToString": {"format": "%Y-%m", "date": "$applied_at"}}, "n": {"$sum": 1}}},
            {"$sort": {"_id": 1}},
        ]
    ).to_list(length=None)
    by_month = [{"month": r["_id"], "n": r["n"]} for r in by_month_raw if r["_id"]]

    return serialize_mongo(
        {
            "range": {"start": start_dt, "end": end_dt - timedelta(days=1)},
            "totals": totals,
            "kpis": kpis,
            "daily": daily_raw,
            "status": status,
            "funnel": funnel,
            "apps_per_student": apps_per_student,
            "top_students": top_students,
            "top_companies": top_companies,
            "role_categories": role_categories,
            "by_month": by_month,
        }
    )


async def get_admin_student_detail(student_id: str) -> dict:
    """Everything about one student: profile, pipeline stats, every application
    (with company/role/status/links/reason), role mix, and interview reports."""
    db = get_database()
    try:
        oid = to_object_id(student_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid student id")

    app_pipeline = [
        {"$match": {"student_id": oid}},
        {"$sort": {"applied_at": -1, "created_at": -1}},
        {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "co"}},
        {"$unwind": {"path": "$co", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "opportunity_id", "foreignField": "_id", "as": "opp"}},
        {"$unwind": {"path": "$opp", "preserveNullAndEmptyArrays": True}},
        {
            "$project": {
                "status": _STATUS,
                "interested": _INTERESTED,
                "applied_at": 1,
                "resume_link": {"$ifNull": ["$application_details.submitted_resume_url", "$resume_link"]},
                "github_link": {"$ifNull": ["$application_details.github_link", "$github_link"]},
                "project_link": {"$ifNull": ["$application_details.project_link", "$project_link"]},
                "interest_reason": "$application_details.interest_reason",
                "non_interest_reason": "$application_details.non_interest_reason",
                "company": "$co.name",
                "company_id": "$co._id",
                "role": "$opp.role",
                "skills": "$opp.must_have_skills",
                "opportunity_id": "$opp._id",
            }
        },
    ]
    reports_pipeline = [
        {"$match": {"student_id": oid}},
        {"$lookup": {"from": INTERVIEW_SESSIONS, "localField": "session_id", "foreignField": "_id", "as": "sess"}},
        {"$unwind": {"path": "$sess", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": TRANSCRIPTS, "localField": "session_id", "foreignField": "session_id", "as": "tr"}},
        {"$unwind": {"path": "$tr", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "co"}},
        {"$unwind": {"path": "$co", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "opportunity_id", "foreignField": "_id", "as": "opp"}},
        {"$unwind": {"path": "$opp", "preserveNullAndEmptyArrays": True}},
        {
            "$project": {
                "overall": 1,
                "communication": 1,
                "strengths": 1,
                "improvements": 1,
                "skill_ratings": 1,
                "interviewer_feedback": 1,
                "visible_to_student": 1,
                "generated_at": 1,
                "created_at": 1,
                "company": "$co.name",
                "role": "$opp.role",
                "interview_date": 1,
                "sess": "$sess",
                "tr": "$tr",
            }
        },
    ]
    # All three are keyed only by the student id and independent of each other —
    # run them as one parallel batch instead of three serial round trips.
    student, app_rows, raw_reports = await asyncio.gather(
        db[STUDENTS].find_one({"_id": oid}),
        db[APPLICATIONS].aggregate(app_pipeline).to_list(length=None),
        db[INTERVIEW_REPORTS].aggregate(reports_pipeline).to_list(length=None),
    )
    if not student:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Student not found")

    stats = {"responses": 0, "interested": 0, "shortlisted": 0, "selected": 0, "declined": 0}
    role_counts: dict = {}
    applications = []
    for a in app_rows:
        st = a.get("status") or "APPLIED"
        interested = a.get("interested") is not False
        stats["responses"] += 1
        stats["interested" if interested else "declined"] += 1
        if st in _SHORTLISTED_SET:
            stats["shortlisted"] += 1
        if st in _SELECTED_SET:
            stats["selected"] += 1
        cat = categorize_role(a.get("role"), a.get("skills"))
        if interested:
            role_counts[cat] = role_counts.get(cat, 0) + 1
        applications.append(
            {
                "company": a.get("company") or "Unknown",
                "id": a.get("_id"),
                "company_id": a.get("company_id"),
                "opportunity_id": a.get("opportunity_id"),
                "role": a.get("role") or "—",
                "category": cat,
                "status": st,
                "interested": interested,
                "applied_at": a.get("applied_at"),
                "resume_link": a.get("resume_link"),
                "github_link": a.get("github_link"),
                "project_link": a.get("project_link"),
                "interest_reason": a.get("interest_reason"),
                "non_interest_reason": a.get("non_interest_reason"),
            }
        )
    role_breakdown = sorted([{"category": k, "n": v} for k, v in role_counts.items()], key=lambda x: -x["n"])

    reports = []
    for r in raw_reports:
        sess = r.get("sess") or {}
        tr = r.get("tr") or {}
        r["interview_date"] = resolve_interview_date(r, session=sess, transcript=tr)
        r["scheduled_at"] = sess.get("scheduled_at")
        r["meeting_date"] = tr.get("meeting_date")
        r.pop("sess", None)
        r.pop("tr", None)
        reports.append(r)

    reports.sort(
        key=lambda r: (
            _sort_key_for_date(r.get("interview_date")),
            _sort_key_for_date(r.get("generated_at")),
            _sort_key_for_date(r.get("created_at")),
        ),
        reverse=True,
    )

    return serialize_mongo(
        {
            "student": {
                "id": student["_id"],
                "name": student.get("name"),
                "email": student.get("email"),
                "phone": student.get("phone"),
                "college_name": student.get("college_name"),
                "current_city": student.get("current_city"),
                "degree": student.get("degree"),
                "department": student.get("department"),
                "year_of_passing": student.get("year_of_passing"),
                "resume_link": student.get("resume_link"),
                "technical_developer_name": student.get("technical_developer_name"),
                "placed_status": student.get("placed_status"),
                "external_user_id": student.get("external_user_id"),
                "created_at": student.get("created_at"),
            },
            "stats": stats,
            "applications": applications,
            "role_breakdown": role_breakdown,
            "reports": reports,
        }
    )


async def list_pending_sessions() -> list[dict]:
    """Sessions that have a transcript with mapped students but no full RSA yet
    (company_expectations not generated). These are what the 'Generate reports'
    button runs, one at a time, so the per-day AI quota isn't blown in a burst."""
    db = get_database()
    sessions = await db[INTERVIEW_SESSIONS].find(
        {"company_expectations.expectations": None}
    ).to_list(length=None)
    if not sessions:
        return []

    # Batch the lookups instead of 3 queries per session (was N+1, slow on a
    # remote Atlas): one query each for transcripts, companies and opportunities.
    session_ids = [se["_id"] for se in sessions]
    company_ids = list({se.get("company_id") for se in sessions if se.get("company_id")})
    opp_ids = list({se.get("opportunity_id") for se in sessions if se.get("opportunity_id")})

    transcripts, companies, opps = await asyncio.gather(
        db[TRANSCRIPTS].find({"session_id": {"$in": session_ids}}, {"session_id": 1, "speaker_map": 1}).to_list(length=None),
        db[COMPANIES].find({"_id": {"$in": company_ids}}, {"name": 1}).to_list(length=None),
        db[HIRING_OPPORTUNITIES].find({"_id": {"$in": opp_ids}}, {"role": 1}).to_list(length=None),
    )
    tx_by_session = {t["session_id"]: t for t in transcripts}
    co_by_id = {c["_id"]: c for c in companies}
    opp_by_id = {o["_id"]: o for o in opps}

    out: list[dict] = []
    for se in sessions:
        transcript = tx_by_session.get(se["_id"])
        mapped = [
            m for m in ((transcript or {}).get("speaker_map") or [])
            if m.get("role") == "student" and m.get("student_id")
        ]
        if not transcript or not mapped:
            continue  # can't analyse without a transcript + mapped students
        out.append({
            "id": se["_id"],
            "company": (co_by_id.get(se.get("company_id")) or {}).get("name") or "Company",
            "role": (opp_by_id.get(se.get("opportunity_id")) or {}).get("role"),
            "students": len(mapped),
        })
    return serialize_mongo(out)


async def list_admin_reports(*, published: bool | None = None, limit: int = 300) -> list[dict]:
    """All interview reports for the admin reports view: newest first, joined to
    student / company / role, with the full report body (admin sees score too).
    published=True -> only shared, published=False -> only pending."""
    db = get_database()
    match: dict = {}
    if published is True:
        match["visible_to_student"] = True
    elif published is False:
        match["visible_to_student"] = {"$ne": True}
    reports = await db[INTERVIEW_REPORTS].aggregate(
        [
            {"$match": match},
            {"$lookup": {"from": INTERVIEW_SESSIONS, "localField": "session_id", "foreignField": "_id", "as": "sess"}},
            {"$unwind": {"path": "$sess", "preserveNullAndEmptyArrays": True}},
            {"$lookup": {"from": TRANSCRIPTS, "localField": "session_id", "foreignField": "session_id", "as": "tr"}},
            {"$unwind": {"path": "$tr", "preserveNullAndEmptyArrays": True}},
            {"$lookup": {"from": STUDENTS, "localField": "student_id", "foreignField": "_id", "as": "st"}},
            {"$unwind": {"path": "$st", "preserveNullAndEmptyArrays": True}},
            {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "co"}},
            {"$unwind": {"path": "$co", "preserveNullAndEmptyArrays": True}},
            {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "opportunity_id", "foreignField": "_id", "as": "opp"}},
            {"$unwind": {"path": "$opp", "preserveNullAndEmptyArrays": True}},
            {
                "$project": {
                    "overall": 1, "communication": 1, "strengths": 1, "improvements": 1,
                    "skill_ratings": 1, "interviewer_feedback": 1, "answers": 1,
                    "interviewer_satisfaction": 1, "coaching_note": 1,
                    "visible_to_student": 1, "generated_at": 1, "created_at": 1, "application_id": 1,
                    "company_id": 1, "company": "$co.name", "role": "$opp.role",
                    "student_id": 1,
                    "student": {
                        "id": "$st._id",
                        "name": "$st.name",
                        "phone": "$st.phone",
                    },
                    "company_expectations": "$sess.company_expectations",
                    "interview_date": 1,
                    "sess": "$sess",
                    "tr": "$tr",
                }
            },
        ]
    ).to_list(length=None)

    for r in reports:
        sess = r.get("sess") or {}
        tr = r.get("tr") or {}
        r["interview_date"] = resolve_interview_date(r, session=sess, transcript=tr)
        r["scheduled_at"] = sess.get("scheduled_at")
        r["meeting_date"] = tr.get("meeting_date")
        r.pop("sess", None)
        r.pop("tr", None)

    reports.sort(
        key=lambda r: (
            _sort_key_for_date(r.get("interview_date")),
            _sort_key_for_date(r.get("generated_at")),
            _sort_key_for_date(r.get("created_at")),
        ),
        reverse=True,
    )
    if limit:
        reports = reports[:limit]

    return serialize_mongo(reports)


def _docx_text(value: object) -> str:
    """Render existing report values without inventing or summarising content."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _docx_list(value: object) -> list[object]:
    if isinstance(value, list):
        return value
    return [value] if value else []


def _add_docx_label(doc: Document, label: str, value: object) -> None:
    text = _docx_text(value)
    if not text:
        return
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(text)


def _add_docx_heading(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text.upper(), level=level)


async def build_company_feedback_docx(company_id: str, month: str | None = None) -> tuple[bytes, str]:
    """Create a DOCX from the same raw report/session fields shown to admins."""
    try:
        company_object_id = to_object_id(company_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid company id")

    db = get_database()
    company = await db[COMPANIES].find_one({"_id": company_object_id}, {"name": 1})
    if not company:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Company not found")

    pipeline: list[dict] = [
        {"$match": {"company_id": company_object_id}},
        {"$lookup": {"from": INTERVIEW_SESSIONS, "localField": "session_id", "foreignField": "_id", "as": "session"}},
        {"$unwind": {"path": "$session", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": TRANSCRIPTS, "localField": "session_id", "foreignField": "session_id", "as": "transcript"}},
        {"$unwind": {"path": "$transcript", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": STUDENTS, "localField": "student_id", "foreignField": "_id", "as": "student"}},
        {"$unwind": {"path": "$student", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "opportunity_id", "foreignField": "_id", "as": "opportunity"}},
        {"$unwind": {"path": "$opportunity", "preserveNullAndEmptyArrays": True}},
    ]

    all_reports = await db[INTERVIEW_REPORTS].aggregate(pipeline).to_list(length=None)

    for r in all_reports:
        sess = r.get("session") or {}
        tr = r.get("transcript") or {}
        r["interview_date"] = resolve_interview_date(r, session=sess, transcript=tr)

    all_reports.sort(
        key=lambda r: (
            _sort_key_for_date(r.get("interview_date")),
            _sort_key_for_date(r.get("generated_at")),
            _sort_key_for_date(r.get("created_at")),
        ),
        reverse=True,
    )

    if month:
        try:
            start = datetime.strptime(month, "%Y-%m").replace(tzinfo=timezone.utc)
        except ValueError:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Month must use YYYY-MM format")
        end = datetime(start.year + (start.month == 12), (start.month % 12) + 1, 1, tzinfo=timezone.utc)

        reports = []
        for r in all_reports:
            idate = r.get("interview_date")
            dt = _sort_key_for_date(idate)
            if dt != datetime.min.replace(tzinfo=timezone.utc) and start <= dt < end:
                reports.append(r)
    else:
        reports = all_reports

    if not reports:
        scope = f" for {month}" if month else ""
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"No interview reports found for this company{scope}")

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10)

    company_name = company.get("name") or "Company"
    document.add_heading("Company Interview Feedback", level=0)
    document.add_heading(company_name, level=1)
    _add_docx_label(document, "Candidates", len(reports))

    company_summary = next((
        (report.get("session") or {}).get("company_expectations")
        for report in reports
        if (report.get("session") or {}).get("company_expectations")
    ), {}) or {}
    if company_summary.get("expectations") or company_summary.get("focus"):
        _add_docx_heading(document, "What This Company Looked For")
        if company_summary.get("expectations"):
            document.add_paragraph(_docx_text(company_summary["expectations"]))
        focus = [_docx_text(item) for item in _docx_list(company_summary.get("focus")) if _docx_text(item)]
        if focus:
            paragraph = document.add_paragraph()
            paragraph.add_run("Requirements: ").bold = True
            paragraph.add_run(" • ".join(focus))

    for index, report in enumerate(reports, start=1):
        if index > 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        document.add_heading(f"Candidate {index}", level=1)
        _add_docx_label(document, "Name", (report.get("student") or {}).get("name") or "Student")
        _add_docx_label(document, "Company", company_name)
        _add_docx_label(document, "Role", (report.get("opportunity") or {}).get("role"))
        score = (report.get("overall") or {}).get("score")
        if score is not None:
            _add_docx_label(document, "Score", f"{score}/10")
        interview_date = resolve_interview_date(report, session=report.get("session"), transcript=report.get("transcript"))
        if interview_date:
            date_text = f"{interview_date.month}/{interview_date.day}/{interview_date.year}" if isinstance(interview_date, datetime) else str(interview_date)
            _add_docx_label(document, "Interview Date", date_text)
        _add_docx_label(document, "Status", "Published" if report.get("visible_to_student") else "Pending")

        overall = report.get("overall") or {}
        if overall.get("summary"):
            _add_docx_heading(document, "Overall Feedback")
            document.add_paragraph(_docx_text(overall["summary"]))
        if report.get("interviewer_satisfaction"):
            _add_docx_heading(document, "How They Met The Bar")
            document.add_paragraph(_docx_text(report["interviewer_satisfaction"]))

        answers = _docx_list(report.get("answers"))
        if answers:
            _add_docx_heading(document, "Questions & Answers")
            for question_index, answer in enumerate(answers, start=1):
                if not isinstance(answer, dict):
                    continue
                document.add_heading(f"Q{question_index}. {_docx_text(answer.get('question_text'))}", level=3)
                accuracy = answer.get("accuracy")
                if accuracy is not None:
                    _add_docx_label(document, "Score", f"{round(float(accuracy) / 20, 1)}/5")
                _add_docx_label(document, "Candidate", answer.get("student_answer"))
                _add_docx_label(document, "Expected", answer.get("ideal_answer"))
                _add_docx_label(document, "Feedback", answer.get("feedback"))

        for heading, key in (("Strengths", "strengths"), ("Areas To Improve", "improvements")):
            items = _docx_list(report.get(key))
            if not items:
                continue
            _add_docx_heading(document, heading)
            for item in items:
                if isinstance(item, dict):
                    content = " — ".join(_docx_text(item.get(field)) for field in ("area", "detail") if _docx_text(item.get(field)))
                else:
                    content = _docx_text(item)
                if content:
                    document.add_paragraph(content, style="List Bullet")

        if report.get("coaching_note"):
            _add_docx_heading(document, "Coaching For Next Time")
            document.add_paragraph(_docx_text(report["coaching_note"]))
        if report.get("interviewer_feedback"):
            _add_docx_heading(document, "Interviewer Feedback")
            document.add_paragraph(_docx_text(report["interviewer_feedback"]))
        communication = report.get("communication") or {}
        if communication.get("notes"):
            _add_docx_heading(document, "Communication")
            document.add_paragraph(_docx_text(communication["notes"]))
        skill_ratings = report.get("skill_ratings") or {}
        if skill_ratings:
            _add_docx_heading(document, "Skill Ratings")
            for skill, rating in skill_ratings.items():
                _add_docx_label(document, str(skill).replace("_", " ").title(), f"{rating}/5")

    buffer = BytesIO()
    document.save(buffer)
    safe_company_name = re.sub(r"[^A-Za-z0-9]+", "_", company_name).strip("_") or "Company"
    month_suffix = f"_{month}" if month else ""
    return buffer.getvalue(), f"{safe_company_name}_Interview_Feedback{month_suffix}.docx"


def _canonical_student_id(report: dict) -> str | None:
    """The canonical identity used across all interview report aggregation."""
    if report.get("student_id") is not None:
        return str(report["student_id"])
    student = report.get("student") or {}
    if isinstance(student, dict):
        value = student.get("id") or student.get("_id")
        if value is not None:
            return str(value)
    return None


def _canonical_company_key(report: dict) -> str | None:
    """Canonical company identity; prefer company_id when available, otherwise use company name as a fallback."""
    if report.get("company_id") is not None:
        return str(report["company_id"])
    company = report.get("company") or {}
    if isinstance(company, dict):
        value = company.get("id") or company.get("_id")
        if value is not None:
            return str(value)
    company_name = company.get("name") if isinstance(company, dict) else report.get("company")
    if company_name:
        return str(company_name)
    return None


def group_reports_by_student_id(reports: list[dict]) -> list[dict]:
    """Group interview reports by canonical student_id without name-based deduplication."""
    grouped: dict[str, dict] = {}
    for report in reports:
        student_id = _canonical_student_id(report)
        if not student_id:
            continue
        student = report.get("student") or {}
        entry = grouped.setdefault(
            student_id,
            {
                "student_id": student_id,
                "student_name": student.get("name") or "Student",
                "student_phone": student.get("phone") or student.get("mobile") or "",
                "reports": [],
                "company_keys": set(),
            },
        )
        if not entry["student_phone"]:
            phone = (student.get("phone") or student.get("mobile") or "")
            entry["student_phone"] = phone
        entry["reports"].append(report)
        company_key = _canonical_company_key(report)
        if company_key is not None:
            entry["company_keys"].add(company_key)

    result = []
    for group in grouped.values():
        result.append(
            {
                "student_id": group["student_id"],
                "student_name": group["student_name"],
                "student_phone": group["student_phone"],
                "reports": group["reports"],
                "company_count": len(group["company_keys"]),
                "interview_count": len(group["reports"]),
            }
        )
    result.sort(key=lambda item: (item["student_name"] or "").lower())
    return result


def resolve_student_export_scope(reports: list[dict], *, scope: str, student_id: str | None = None) -> dict:
    """Resolve one clicked report or a student's entire interview history using student_id."""
    if not reports:
        raise ValueError("No interview reports were provided for export")

    if scope == "single":
        report = reports[0]
        resolved_student_id = student_id or _canonical_student_id(report)
        if not resolved_student_id:
            raise ValueError("This report has no valid student_id; student history cannot be safely aggregated.")
        return {"student_id": str(resolved_student_id), "reports": [report]}

    if scope == "student":
        resolved_student_id = student_id
        if not resolved_student_id:
            raise ValueError("student_id is required to aggregate a student's interview history")
        filtered = [
            report for report in reports
            if (_canonical_student_id(report) == str(resolved_student_id)) or (report.get("student_id") is not None and str(report["student_id"]) == str(resolved_student_id))
        ]
        if not filtered:
            raise ValueError("No interview reports were found for this student_id")
        return {"student_id": str(resolved_student_id), "reports": filtered}

    raise ValueError(f"Unsupported export scope: {scope}")


async def resolve_student_report_ids(student_id: str) -> list[str]:
    """Return all report ids for a canonical student_id without using name-based grouping."""
    return await resolve_student_ids_to_report_ids([student_id])


async def resolve_student_ids_to_report_ids(student_ids: list[str]) -> list[str]:
    """Resolve multiple canonical student ids to all matching report ids without using names."""
    clean_ids = [student_id for student_id in student_ids if student_id]
    if not clean_ids:
        return []

    object_ids: list[ObjectId] = []
    for student_id in clean_ids:
        try:
            object_ids.append(to_object_id(student_id))
        except ValueError as exc:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid student id") from exc

    db = get_database()
    reports = await db[INTERVIEW_REPORTS].find({"student_id": {"$in": object_ids}}, {"_id": 1}).sort([("generated_at", -1), ("created_at", -1)]).to_list(length=None)
    return [str(report["_id"]) for report in reports]


def _sanitize_filename(text: str) -> str:
    """Sanitize text for safe use in filenames while preserving readability."""
    if not text:
        return "Item"
    safe = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return safe or "Item"


def _student_filename(report: dict, used_names: set[str] | None = None) -> str:
    """Generate filename for a single student interview: StudentName_CompanyName_Interview_Feedback.docx"""
    name = (report.get("student") or {}).get("name") or "Student"
    safe_name = _sanitize_filename(name)
    company_name = (report.get("company") or {}).get("name") or "Interview"
    safe_company = _sanitize_filename(company_name)
    filename = f"{safe_name}_{safe_company}_Interview_Feedback.docx"
    if used_names is not None and filename in used_names:
        suffix = str(report.get("_id") or report.get("id") or "")[-6:]
        filename = f"{safe_name}_{safe_company}_{suffix}_Interview_Feedback.docx"
    if used_names is not None:
        used_names.add(filename)
    return filename


def _export_filename_for_scope(
    student_name: str,
    company_names: list[str],
    scope: str,
    total_student_companies: int,
) -> str:
    """Generate the combined DOCX/ZIP filename based on export scope.
    
    - scope="single": StudentName_CompanyName_Interview_Feedback.docx
    - scope="student" (all companies): StudentName_Interview_Feedback.docx
    - scope="selected" (subset): StudentName_Selected_Companies_Interview_Feedback.docx
    """
    safe_name = _sanitize_filename(student_name)
    
    if scope == "single" and company_names:
        # Single interview: include the company name
        safe_company = _sanitize_filename(company_names[0])
        return f"{safe_name}_{safe_company}_Interview_Feedback.docx"
    elif scope == "student":
        # All companies for this student
        return f"{safe_name}_Interview_Feedback.docx"
    elif scope == "selected" and len(company_names) < total_student_companies:
        # Multiple but not all companies
        return f"{safe_name}_Selected_Companies_Interview_Feedback.docx"
    else:
        # Fallback for "selected" that happens to be all companies
        return f"{safe_name}_Interview_Feedback.docx"


def _new_feedback_document(title: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading(title, level=0)
    return document


def _add_student_feedback_section(document: Document, report: dict, index: int, *, page_break: bool = False) -> None:
    """Write the interview-feedback content only; no extra metadata markers or numbering."""
    if page_break:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    company_name = (report.get("company") or {}).get("name") or "Company"
    document.add_heading(company_name, level=1)
    _add_docx_label(document, "Role", (report.get("opportunity") or {}).get("role"))

    interview_date = resolve_interview_date(report, session=report.get("session"), transcript=report.get("transcript"))
    if interview_date:
        date_text = f"{interview_date.month}/{interview_date.day}/{interview_date.year}" if isinstance(interview_date, datetime) else str(interview_date)
        _add_docx_label(document, "Interview Date", date_text)

    score = (report.get("overall") or {}).get("score")
    if score is not None:
        _add_docx_label(document, "Overall Score", f"{score}/10")

    overall = report.get("overall") or {}
    if overall.get("summary"):
        _add_docx_heading(document, "Overall Feedback")
        document.add_paragraph(_docx_text(overall["summary"]))
    if report.get("interviewer_satisfaction"):
        _add_docx_heading(document, "How They Met The Bar")
        document.add_paragraph(_docx_text(report["interviewer_satisfaction"]))

    answers = _docx_list(report.get("answers"))
    if answers:
        _add_docx_heading(document, "Questions & Answers")
        for question_index, answer in enumerate(answers, start=1):
            if not isinstance(answer, dict):
                continue
            document.add_heading(f"Q{question_index}. {_docx_text(answer.get('question_text'))}", level=3)
            _add_docx_label(document, "Candidate Answer", answer.get("student_answer"))
            accuracy = answer.get("accuracy")
            if accuracy is not None:
                _add_docx_label(document, "Score", f"{round(float(accuracy) / 20, 1)}/5")
            _add_docx_label(document, "Expected Answer", answer.get("ideal_answer"))
            _add_docx_label(document, "Feedback", answer.get("feedback"))

    for heading, key in (("Strengths", "strengths"), ("Areas To Improve", "improvements")):
        items = _docx_list(report.get(key))
        if items:
            _add_docx_heading(document, heading)
            for item in items:
                content = " — ".join(_docx_text(item.get(field)) for field in ("area", "detail") if _docx_text(item.get(field))) if isinstance(item, dict) else _docx_text(item)
                if content:
                    document.add_paragraph(content, style="List Bullet")

    communication = report.get("communication") or {}
    if communication.get("notes"):
        _add_docx_heading(document, "Communication")
        document.add_paragraph(_docx_text(communication["notes"]))

    skill_ratings = report.get("skill_ratings") or {}
    if skill_ratings:
        _add_docx_heading(document, "Skill Ratings")
        for skill, rating in skill_ratings.items():
            _add_docx_label(document, str(skill).replace("_", " ").title(), f"{rating}/5")


async def build_student_feedback_export(
    report_ids: list[str],
    mode: str,
    scope: str | None = None,
    student_id: str | None = None,
) -> tuple[bytes, str, str]:
    """Build a DOCX or ZIP from selected existing report documents, in UI order.

    For multi-student exports, package all selected students into one ZIP so the
    browser only receives a single download.
    """
    try:
        object_ids = [to_object_id(report_id) for report_id in report_ids]
    except ValueError:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid report id")

    db = get_database()
    reports = await db[INTERVIEW_REPORTS].aggregate([
        {"$match": {"_id": {"$in": object_ids}}},
        {"$lookup": {"from": STUDENTS, "localField": "student_id", "foreignField": "_id", "as": "student"}},
        {"$unwind": {"path": "$student", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "company"}},
        {"$unwind": {"path": "$company", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "opportunity_id", "foreignField": "_id", "as": "opportunity"}},
        {"$unwind": {"path": "$opportunity", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": INTERVIEW_SESSIONS, "localField": "session_id", "foreignField": "_id", "as": "session"}},
        {"$unwind": {"path": "$session", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": TRANSCRIPTS, "localField": "session_id", "foreignField": "session_id", "as": "transcript"}},
        {"$unwind": {"path": "$transcript", "preserveNullAndEmptyArrays": True}},
    ]).to_list(length=None)

    by_id = {str(report["_id"]): report for report in reports}
    ordered_reports = [by_id[str(report_id)] for report_id in object_ids if str(report_id) in by_id]
    if not ordered_reports:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No selected interview reports were found")

    grouped_by_student = group_reports_by_student_id(ordered_reports)
    student_count = len(grouped_by_student)
    first_student_name = (ordered_reports[0].get("student") or {}).get("name") or "Student"

    def save_document(document: Document) -> bytes:
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def build_student_doc(report_list: list[dict]) -> bytes:
        student_name = (report_list[0].get("student") or {}).get("name") or "Student"
        student_doc = _new_feedback_document("STUDENT INTERVIEW FEEDBACK")
        student_doc.add_paragraph(f"Student: {student_name}")
        student_doc.add_paragraph(f"Total Interviews: {len(report_list)}")
        student_doc.add_paragraph()
        student_doc.add_heading("COMPANIES INTERVIEWED", level=1)
        for index, report in enumerate(report_list, start=1):
            company_name = (report.get("company") or {}).get("name") or "Company"
            student_doc.add_paragraph(f"{index}. {company_name}")
        for index, report in enumerate(report_list, start=1):
            _add_student_feedback_section(student_doc, report, index, page_break=True)
        return save_document(student_doc)

    if mode == "combined":
        if student_count > 1:
            return build_student_doc(ordered_reports), "Student_Interview_Feedback_Combined.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        selected_companies = list(dict.fromkeys((report.get("company") or {}).get("name") or "Company" for report in ordered_reports))
        combined_name = _export_filename_for_scope(first_student_name, selected_companies, scope or "selected", len(selected_companies))
        return build_student_doc(ordered_reports), combined_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    archive = BytesIO()
    used_names: set[str] = set()
    with zipfile.ZipFile(archive, "w", zipfile.ZIP_DEFLATED) as zip_file:
        if mode == "both":
            if student_count > 1:
                zip_file.writestr("Student_Interview_Feedback_Combined.docx", build_student_doc(ordered_reports))
            else:
                selected_companies = list(dict.fromkeys((report.get("company") or {}).get("name") or "Company" for report in ordered_reports))
                combined_name = _export_filename_for_scope(first_student_name, selected_companies, scope or "selected", len(selected_companies))
                zip_file.writestr(combined_name, build_student_doc(ordered_reports))

        if student_count > 1:
            for group in grouped_by_student:
                student_name = group["student_name"] or "Student"
                company_names = list(dict.fromkeys((report.get("company") or {}).get("name") or "Company" for report in group["reports"]))
                filename = _export_filename_for_scope(student_name, company_names, "student", len(company_names))
                zip_file.writestr(filename, build_student_doc(group["reports"]))
        else:
            for report in ordered_reports:
                student_doc = _new_feedback_document("STUDENT INTERVIEW FEEDBACK")
                student_name = (report.get("student") or {}).get("name") or "Student"
                student_doc.add_paragraph(f"Student: {student_name}")
                student_doc.add_paragraph(f"Total Interviews: {len(ordered_reports)}")
                student_doc.add_paragraph()
                student_doc.add_heading("COMPANIES INTERVIEWED", level=1)
                company_name = (report.get("company") or {}).get("name") or "Company"
                student_doc.add_paragraph(f"1. {company_name}")
                _add_student_feedback_section(student_doc, report, 1, page_break=True)
                filename = _student_filename(report, used_names)
                zip_file.writestr(filename, save_document(student_doc))

    zip_filename = "Student_Interview_Feedback.zip" if student_count > 1 else f"{_sanitize_filename(first_student_name)}_Interview_Feedback.zip"
    return archive.getvalue(), zip_filename, "application/zip"


async def build_company_feedback_export(report_ids: list[str], mode: str) -> tuple[bytes, str, str]:
    """Build a company-oriented DOCX/ZIP from the report ids selected by UI filters."""
    try:
        object_ids = [to_object_id(report_id) for report_id in report_ids]
    except ValueError:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Invalid report id")
    db = get_database()
    reports = await db[INTERVIEW_REPORTS].aggregate([
        {"$match": {"_id": {"$in": object_ids}}},
        {"$lookup": {"from": STUDENTS, "localField": "student_id", "foreignField": "_id", "as": "student"}},
        {"$unwind": {"path": "$student", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": COMPANIES, "localField": "company_id", "foreignField": "_id", "as": "company"}},
        {"$unwind": {"path": "$company", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": HIRING_OPPORTUNITIES, "localField": "opportunity_id", "foreignField": "_id", "as": "opportunity"}},
        {"$unwind": {"path": "$opportunity", "preserveNullAndEmptyArrays": True}},
        {"$lookup": {"from": INTERVIEW_SESSIONS, "localField": "session_id", "foreignField": "_id", "as": "session"}},
        {"$unwind": {"path": "$session", "preserveNullAndEmptyArrays": True}},
    ]).to_list(length=None)
    by_id = {str(report["_id"]): report for report in reports}
    ordered_reports = [by_id[str(report_id)] for report_id in object_ids if str(report_id) in by_id]
    if not ordered_reports:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No selected interview reports were found")

    companies: dict[str, list[dict]] = {}
    for report in ordered_reports:
        name = (report.get("company") or {}).get("name") or "Company"
        companies.setdefault(name, []).append(report)

    def add_company_section(document: Document, name: str, company_reports: list[dict], *, page_break: bool = False) -> None:
        if page_break:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        document.add_heading(name, level=1)
        _add_docx_label(document, "Candidates", len(company_reports))
        expectations = next(((report.get("session") or {}).get("company_expectations") for report in company_reports if (report.get("session") or {}).get("company_expectations")), {}) or {}
        if expectations.get("expectations") or expectations.get("focus"):
            _add_docx_heading(document, "What This Company Looked For")
            if expectations.get("expectations"):
                document.add_paragraph(_docx_text(expectations["expectations"]))
            focus = [_docx_text(item) for item in _docx_list(expectations.get("focus")) if _docx_text(item)]
            if focus:
                _add_docx_label(document, "Requirement tags", " • ".join(focus))
        _add_docx_heading(document, "Candidates")
        for index, report in enumerate(company_reports, start=1):
            _add_student_feedback_section(document, report, index)

    def save_document(document: Document) -> bytes:
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    combined = None
    if mode in {"combined", "both"}:
        combined_doc = _new_feedback_document("Company Interview Feedback")
        _add_docx_label(combined_doc, "Generated", datetime.now(timezone.utc).strftime("%d %B %Y"))
        _add_docx_label(combined_doc, "Companies included", len(companies))
        _add_docx_heading(combined_doc, "Table of Contents")
        for index, name in enumerate(companies, start=1):
            combined_doc.add_paragraph(f"{index}. {name}")
        for index, (name, company_reports) in enumerate(companies.items(), start=1):
            add_company_section(combined_doc, name, company_reports, page_break=True)
        combined = save_document(combined_doc)
        if mode == "combined":
            if len(companies) == 1:
                safe_name = re.sub(r"[^A-Za-z0-9]+", "_", next(iter(companies))).strip("_") or "Company"
                return combined, f"{safe_name}_Interview_Feedback.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            return combined, "Company_Feedback_Combined.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # One company never needs a ZIP: separate and combined output would be the
    # same company-level document, so return a single clean DOCX instead.
    if len(companies) == 1:
        name, company_reports = next(iter(companies.items()))
        company_doc = _new_feedback_document("Company Interview Feedback")
        add_company_section(company_doc, name, company_reports)
        safe_name = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_") or "Company"
        return save_document(company_doc), f"{safe_name}_Interview_Feedback.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    archive = BytesIO()
    used_names: set[str] = set()
    with zipfile.ZipFile(archive, "w", zipfile.ZIP_DEFLATED) as zip_file:
        if combined:
            zip_file.writestr("Company_Feedback_Combined.docx", combined)
        for name, company_reports in companies.items():
            company_doc = _new_feedback_document("Company Interview Feedback")
            add_company_section(company_doc, name, company_reports)
            safe_name = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_") or "Company"
            filename = f"{safe_name}_Interview_Feedback.docx"
            if filename in used_names:
                filename = f"{safe_name}_{str(company_reports[0].get('_id'))[-6:]}_Interview_Feedback.docx"
            used_names.add(filename)
            zip_file.writestr(filename, save_document(company_doc))
    return archive.getvalue(), "Company_Feedback_Selected.zip", "application/zip"
